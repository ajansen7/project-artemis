#!/usr/bin/env python3
"""
Resume DOCX/PDF Generator — Builds a styled resume from markdown.
Clean modern header (no table), hyperlinked contact line, two-line role entries,
light-blue summary box, ALL-CAPS blue section headings, 2-col skills table.

Usage:
  uv run python tools/generate_resume_docx.py --job-id <uuid>
  uv run python tools/generate_resume_docx.py --resume-path path/to/resume.md

Requires LibreOffice for PDF conversion:
  brew install --cask libreoffice
"""

import argparse
import os
import re
import subprocess
import sys
from pathlib import Path

from dotenv import load_dotenv

PROJECT_ROOT = Path(__file__).resolve().parent.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))
load_dotenv(PROJECT_ROOT / ".env")

SUPABASE_URL = os.getenv("SUPABASE_URL", "")
SUPABASE_KEY = os.getenv("SUPABASE_SERVICE_ROLE_KEY", "")

# ─── Colour palette ───────────────────────────────────────────────────────────

BLUE   = "2E7EBF"   # section headings + summary border + accent rule
SUMMBG = "EAF3FA"   # summary box background
DARK   = "1A1A2E"   # name / body text (near-black)
GREY   = "646464"   # subtitle / dates / contact line


# ─── Markdown parser ──────────────────────────────────────────────────────────

def parse_resume_md(md_text: str) -> list[tuple[str, str]]:
    """
    Parse resume markdown into (block_type, content) tuples.

    Types: name, contact, section, role, bullet, hr, blank, paragraph
    """
    blocks: list[tuple[str, str]] = []
    seen_name = False
    seen_contact = False

    for raw in md_text.split("\n"):
        line = raw.rstrip()
        if not line:
            blocks.append(("blank", ""))
            continue

        if line.startswith("# "):
            t, content = "name", line[2:].strip()
            seen_name = True
        elif line.startswith("## "):
            t, content = "section", line[3:].strip()
        elif line.startswith("### "):
            t, content = "role", line[4:].strip()
        elif line.startswith("- ") or line.startswith("* "):
            t, content = "bullet", line[2:].strip()
        elif line.strip() in ("---", "***", "___"):
            t, content = "hr", ""
        elif seen_name and not seen_contact and not line.startswith("#"):
            t = "contact"
            content = re.sub(r"^\*\*Contact\*\*:\s*", "", line.strip())
            seen_contact = True
        elif re.match(r"^\*\*[^*]+\*\*", line.strip()):
            # Sub-role header — may have trailing italic parenthetical or plain text.
            # The leading **...** is always treated as bold in _sub_role, even if the
            # inline parser wouldn't pick it up (e.g., whole-line bold like "**Foo**").
            t, content = "subrole", line.strip()
        elif re.match(r"^\*[^*].+\*$", line.strip()):
            t, content = "italic", line.strip()[1:-1]
        else:
            t, content = "paragraph", line.strip()

        blocks.append((t, content))

    return blocks


def parse_role_line(text: str) -> tuple[str, str, str]:
    """Parse '### Company — Title | dates' → (title, company, dates)."""
    pipe_match = re.search(r"\s+\|\s+(.+)$", text)
    if pipe_match:
        dates = pipe_match.group(1).strip()
        remainder = text[: pipe_match.start()].strip()
    else:
        paren_match = re.search(r"\(([^()]+)\)\s*$", text)
        if paren_match:
            dates = paren_match.group(1).strip()
            remainder = text[: paren_match.start()].strip()
        else:
            dates, remainder = "", text

    for sep in (" \u2014 ", " \u2013 ", " — ", " – ", " - "):
        if sep in remainder:
            parts = remainder.split(sep, 1)
            return parts[1].strip(), parts[0].strip(), dates

    return remainder, "", dates


def parse_contact_items(contact_line: str) -> list[dict]:
    """Split 'email | phone | url' into [{'label': str, 'url': str|None}]."""
    items = []
    for part in contact_line.split(" | "):
        part = part.strip()
        m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
        if m:
            items.append({"label": m.group(1), "url": m.group(2)})
        else:
            items.append({"label": part, "url": None})
    return items


# Safety net: strips a trailing run of [tag] markers (e.g., "... automation [ops] [0to1]")
# from resume bullets. The apply skill is responsible for never emitting these in the first
# place — this is belt-and-suspenders for stale tailored resumes.
_TRAILING_TAGS = re.compile(r"(?:\s*\[[A-Za-z0-9_-]+\])+\s*$")


def _parse_inline_runs(text: str) -> list[dict]:
    """Parse **bold**, *italic*, [link](url) into run dicts."""
    runs = []
    pattern = re.compile(
        r"\*\*(.+?)\*\*|\*(.+?)\*|__(.+?)__|_(.+?)_|\[([^\]]+)\]\(([^)]+)\)"
    )
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            runs.append({"text": text[pos : m.start()], "bold": False, "italic": False, "url": None})
        if m.group(1):
            runs.append({"text": m.group(1), "bold": True,  "italic": False, "url": None})
        elif m.group(2):
            runs.append({"text": m.group(2), "bold": False, "italic": True,  "url": None})
        elif m.group(3):
            runs.append({"text": m.group(3), "bold": True,  "italic": False, "url": None})
        elif m.group(4):
            runs.append({"text": m.group(4), "bold": False, "italic": True,  "url": None})
        elif m.group(5):
            runs.append({"text": m.group(5), "bold": False, "italic": False, "url": m.group(6)})
        pos = m.end()
    if pos < len(text):
        runs.append({"text": text[pos:], "bold": False, "italic": False, "url": None})
    return runs or [{"text": text, "bold": False, "italic": False, "url": None}]


# ─── Low-level helpers ────────────────────────────────────────────────────────

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _rgb(hex_str: str) -> RGBColor:
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _font(run, size=None, bold=None, italic=None, color: str | None = None):
    run.font.name = "Calibri"
    if size   is not None: run.font.size  = Pt(size)
    if bold   is not None: run.font.bold  = bold
    if italic is not None: run.font.italic = italic
    if color  is not None: run.font.color.rgb = _rgb(color)


def _spacing(p, before=0, after=0, line=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line is not None:
        pf.line_spacing = Pt(line)


def _shd(cell, fill: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"),  "clear")
    tcPr.append(shd)


def _cell_margins(cell, top=0, left=0, bottom=0, right=0):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("left", left), ("bottom", bottom), ("right", right)]:
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:type"), "dxa")
        e.set(qn("w:w"), str(val))
        tcMar.append(e)
    tcPr.append(tcMar)


def _no_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBdr = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        tcBdr.append(b)
    tcPr.append(tcBdr)


def _get_or_add_tblPr(tbl_elem):
    tblPr = tbl_elem.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_elem.insert(0, tblPr)
    return tblPr


def _no_tbl_borders(table):
    tblPr = _get_or_add_tblPr(table._tbl)
    tblBdr = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        tblBdr.append(b)
    tblPr.append(tblBdr)


def _valign(cell, val="center"):
    tcPr = cell._tc.get_or_add_tcPr()
    v = OxmlElement("w:vAlign")
    v.set(qn("w:val"), val)
    tcPr.append(v)


def _set_col_widths(table, widths: list[int]):
    """Set column widths in twips."""
    tbl = table._tbl
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid")
        tbl.append(tblGrid)
    for gc in tblGrid.findall(qn("w:gridCol")):
        tblGrid.remove(gc)
    for w in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tblGrid.append(gc)
    for cell, w in zip(table.rows[0].cells, widths):
        tcPr = cell._tc.get_or_add_tcPr()
        existing = tcPr.find(qn("w:tcW"))
        if existing is not None:
            tcPr.remove(existing)
        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:type"), "dxa")
        tcW.set(qn("w:w"), str(w))
        tcPr.append(tcW)


def _set_tbl_width(table, width: int):
    tblPr = _get_or_add_tblPr(table._tbl)
    existing = tblPr.find(qn("w:tblW"))
    if existing is not None:
        tblPr.remove(existing)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(width))
    tblPr.append(tblW)


def _inline_runs(para, text: str, size=10, color: str | None = None):
    """Add inline-markup-aware runs to a paragraph."""
    for rd in _parse_inline_runs(text):
        run = para.add_run(rd["text"])
        _font(run, size=size, bold=rd["bold"], italic=rd["italic"],
              color=color or DARK)


def _add_hyperlink(para, text: str, url: str, size: float = 9, color: str = GREY):
    """Inject a w:hyperlink run into para with proper relationship and styling."""
    part = para.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    clr = OxmlElement("w:color")
    clr.set(qn("w:val"), color)
    rPr.append(clr)

    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hyperlink.append(r)
    para._p.append(hyperlink)


# ─── Section builders ─────────────────────────────────────────────────────────

def _header_block(doc, name: str, subtitle: str, contact_items: list[dict]):
    """Clean prose header: name, subtitle, contact line (with blue accent rule baked into its bottom border)."""
    # Name — tight line-height so the subtitle sits directly below
    p_name = doc.add_paragraph()
    _spacing(p_name, before=0, after=0, line=24)
    _font(p_name.add_run(name), size=22, bold=True, color=DARK)

    # Subtitle (role/tagline) — pulled up under the name
    if subtitle:
        p_sub = doc.add_paragraph()
        _spacing(p_sub, before=0, after=2, line=13)
        _font(p_sub.add_run(subtitle), size=11, color=GREY)

    # Contact line — items joined by  ·  separator. Blue accent rule attached directly
    # as this paragraph's bottom border so there's no empty-paragraph gap above the rule.
    if contact_items:
        p_contact = doc.add_paragraph()
        _spacing(p_contact, before=0, after=6, line=13)
        first = True
        for item in contact_items:
            if not item["label"]:
                continue
            if not first:
                sep = p_contact.add_run("  ·  ")
                _font(sep, size=9, color=GREY)
            first = False
            if item.get("url"):
                _add_hyperlink(p_contact, item["label"], item["url"], size=9, color=GREY)
            else:
                run = p_contact.add_run(item["label"])
                _font(run, size=9, color=GREY)

        pPr = p_contact._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), BLUE)
        pBdr.append(bottom)
        pPr.append(pBdr)


def _summary_box(doc, text: str):
    """Full-width light-blue box with thick blue left border."""
    table = doc.add_table(rows=1, cols=1)
    _set_tbl_width(table, 10080)
    _set_col_widths(table, [10080])

    # Blue left border only
    tblPr = _get_or_add_tblPr(table._tbl)
    tblBdr = OxmlElement("w:tblBorders")
    for side in ["top", "bottom", "right", "insideH", "insideV"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        tblBdr.append(b)
    left_bdr = OxmlElement("w:left")
    left_bdr.set(qn("w:val"), "single")
    left_bdr.set(qn("w:color"), BLUE)
    left_bdr.set(qn("w:sz"), "24")
    tblBdr.append(left_bdr)
    tblPr.append(tblBdr)

    cell = table.cell(0, 0)
    _shd(cell, SUMMBG)
    _cell_margins(cell, top=120, left=240, bottom=120, right=240)
    _no_cell_borders(cell)

    p = cell.paragraphs[0]
    _spacing(p, before=6, after=0)
    _inline_runs(p, text, size=10, color=DARK)


def _section_heading(doc, text: str):
    p = doc.add_paragraph()
    _spacing(p, before=10, after=3, line=13)
    _font(p.add_run(text.upper()), size=11, bold=True, color=BLUE)
    # Rule rendered as bottom-border of the heading paragraph itself — keeps the rule
    # tight against the header text instead of pushing it down by an empty paragraph's line-height.
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BBBBBB")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _role_header(doc, title: str, company: str, dates: str):
    if company:
        # Line 1: Company (bold) + right-tab + dates (grey)
        p1 = doc.add_paragraph()
        _spacing(p1, before=10, after=1)
        p1.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), WD_ALIGN_PARAGRAPH.RIGHT)
        _font(p1.add_run(company), size=11, bold=True, color=DARK)
        if dates:
            _font(p1.add_run(f"\t{dates}"), size=9, color=GREY)
        # Line 2: Title (italic, flush left)
        p2 = doc.add_paragraph()
        _spacing(p2, before=0, after=1)
        _font(p2.add_run(title), size=10, italic=True, color=DARK)
    else:
        p = doc.add_paragraph()
        _spacing(p, before=6, after=1)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), WD_ALIGN_PARAGRAPH.RIGHT)
        _font(p.add_run(title), size=11, bold=True, color=DARK)
        if dates:
            _font(p.add_run(f"\t{dates}"), size=9, color=GREY)


def _sub_role(doc, content: str):
    """Secondary heading inside a role block (e.g., "**AI Evaluation Infrastructure** *(parallel role)*").

    Always force-bolds the leading **...** segment so sub-roles render consistently
    even if upstream markdown is inconsistent. Any trailing content after the bold
    segment is rendered with normal inline-markup parsing (picks up italic, etc.).
    """
    p = doc.add_paragraph()
    _spacing(p, before=3, after=1)

    m = re.match(r"^\*\*([^*]+)\*\*\s*(.*)$", content)
    if m:
        _font(p.add_run(m.group(1)), size=10.5, bold=True, color=DARK)
        tail = m.group(2).strip()
        if tail:
            p.add_run(" ")
            for rd in _parse_inline_runs(tail):
                run = p.add_run(rd["text"])
                _font(run, size=10.5, bold=rd["bold"], italic=rd["italic"], color=DARK)
    else:
        # Fallback — line started with `**` but didn't close cleanly; still render bold.
        _font(p.add_run(content), size=10.5, bold=True, color=DARK)


def _bullet(doc, text: str):
    text = _TRAILING_TAGS.sub("", text).rstrip()
    p = doc.add_paragraph()
    _spacing(p, before=1.5, after=1.5, line=13)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    bullet_run = p.add_run("•  ")
    _font(bullet_run, size=10, color=BLUE)
    _inline_runs(p, text, size=10)


def _skills_table(doc, pairs: list[tuple[str, str]]):
    """Two-column table: bold label | value."""
    table = doc.add_table(rows=len(pairs), cols=2)
    _no_tbl_borders(table)
    _set_tbl_width(table, 10080)

    col_widths = [2600, 7480]
    tbl = table._tbl
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid")
        tbl.append(tblGrid)
    for gc in tblGrid.findall(qn("w:gridCol")):
        tblGrid.remove(gc)
    for w in col_widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tblGrid.append(gc)

    for i, (label, value) in enumerate(pairs):
        row = table.rows[i]
        for j, (cell, w) in enumerate(zip(row.cells, col_widths)):
            _no_cell_borders(cell)
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:type"), "dxa")
            tcW.set(qn("w:w"), str(w))
            tcPr.append(tcW)

        lp = row.cells[0].paragraphs[0]
        _spacing(lp, before=1, after=1)
        _font(lp.add_run(f"{label}:"), size=10, bold=True)

        vp = row.cells[1].paragraphs[0]
        _spacing(vp, before=1, after=1)
        _font(vp.add_run(value), size=10)


# ─── Main builder ─────────────────────────────────────────────────────────────

def markdown_to_docx(md_text: str, output_path: str) -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    doc.styles["Normal"].paragraph_format.space_after = Pt(0)

    blocks = parse_resume_md(md_text)

    # Extract header fields
    name         = next((c for t, c in blocks if t == "name"),    "")
    contact_raw  = next((c for t, c in blocks if t == "contact"), "")
    contact_items = parse_contact_items(contact_raw) if contact_raw else []
    first_role   = next((c for t, c in blocks if t == "role"),    None)
    subtitle     = parse_role_line(first_role)[0] if first_role else ""

    _header_block(doc, name, subtitle, contact_items)

    # Find where body starts (skip name/contact/blank/hr preamble)
    i = 0
    while i < len(blocks) and blocks[i][0] in ("name", "contact", "blank", "hr"):
        i += 1

    # Collect summary paragraph(s) before first section or role
    summary_parts: list[str] = []
    j = i
    while j < len(blocks) and blocks[j][0] in ("paragraph", "blank", "hr"):
        if blocks[j][0] == "paragraph":
            summary_parts.append(blocks[j][1])
        j += 1
    if summary_parts:
        _summary_box(doc, " ".join(summary_parts))
        i = j

    # Render body
    current_section = ""
    skill_pairs: list[tuple[str, str]] = []

    def flush_skills():
        if skill_pairs:
            _skills_table(doc, list(skill_pairs))
            skill_pairs.clear()

    while i < len(blocks):
        btype, content = blocks[i]
        i += 1

        if btype in ("blank", "hr"):
            continue

        if btype == "section":
            flush_skills()
            current_section = content.lower()
            # Suppress the "About" heading — the paragraph beneath it is self-explanatory
            # as the top-of-resume summary. Keeps the header block flowing straight into Experience.
            if current_section == "about":
                continue
            _section_heading(doc, content)
            continue

        if btype == "role":
            flush_skills()
            title, company, dates = parse_role_line(content)
            _role_header(doc, title, company, dates)
            continue

        if btype == "bullet":
            if current_section == "skills":
                m = re.match(r"\*\*(.+?)\*\*:\s*(.*)", content)
                if m:
                    skill_pairs.append((m.group(1), m.group(2)))
                    continue
            flush_skills()
            _bullet(doc, content)
            continue

        if btype == "subrole":
            flush_skills()
            _sub_role(doc, content)
            continue

        if btype in ("paragraph", "italic"):
            flush_skills()
            p = doc.add_paragraph()
            _spacing(p, before=2, after=2)
            italic_override = btype == "italic"
            for rd in _parse_inline_runs(content):
                run = p.add_run(rd["text"])
                _font(run, size=10, bold=rd["bold"],
                      italic=rd["italic"] or italic_override)
            continue

    flush_skills()

    docx_out = output_path if output_path.endswith(".docx") else output_path.replace(".pdf", ".docx")
    doc.save(docx_out)
    print(f"✅ DOCX written to: {docx_out}")


# ─── PDF conversion ───────────────────────────────────────────────────────────

def docx_to_pdf(docx_path: str, output_dir: str) -> str:
    soffice = _find_soffice()
    if not soffice:
        print("❌ LibreOffice not found. Install with:")
        print("   brew install --cask libreoffice")
        sys.exit(1)
    result = subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", output_dir, docx_path],
        capture_output=True, text=True, timeout=60,
    )
    if result.returncode != 0:
        print(f"❌ LibreOffice conversion failed:\n{result.stderr}")
        sys.exit(1)
    pdf_path = str(Path(output_dir) / f"{Path(docx_path).stem}.pdf")
    print(f"✅ PDF written to: {pdf_path}")
    return pdf_path


def markdown_to_pdf(md_text: str, output_path: str) -> None:
    docx_path = output_path.replace(".pdf", ".docx")
    markdown_to_docx(md_text, docx_path)
    docx_to_pdf(docx_path, str(Path(output_path).parent))


def _find_soffice() -> str | None:
    for c in [
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "/usr/bin/soffice",
        "/usr/local/bin/soffice",
    ]:
        if Path(c).exists():
            return c
    result = subprocess.run(["which", "soffice"], capture_output=True, text=True)
    return result.stdout.strip() if result.returncode == 0 else None


# ─── DB helpers ───────────────────────────────────────────────────────────────

def fetch_resume_md_from_db(job_id: str) -> tuple[str, str, str]:
    if not SUPABASE_URL or not SUPABASE_KEY:
        print("ERROR: SUPABASE_URL and SUPABASE_SERVICE_ROLE_KEY must be set in .env")
        sys.exit(1)
    from supabase import create_client
    sb = create_client(SUPABASE_URL, SUPABASE_KEY)

    app_res = sb.table("applications").select("resume_md").eq("job_id", job_id).execute()
    if not app_res.data or not app_res.data[0].get("resume_md"):
        print(f"❌ No resume_md found for job {job_id}. Run /generate first.")
        sys.exit(1)

    job_res = sb.table("jobs").select("title, companies(name)").eq("id", job_id).execute()
    if not job_res.data:
        print(f"❌ Job {job_id} not found.")
        sys.exit(1)

    job = job_res.data[0]
    company = (job.get("companies") or {}).get("name", "unknown")
    title   = job.get("title", "unknown")
    return app_res.data[0]["resume_md"], company, title


def _upload_artifact_to_storage(job_id: str, file_path: str, company: str, title: str) -> str:
    """Upload a generated artifact to Supabase Storage, overwriting any existing
    object at the same path. Returns the storage bucket path on success.
    Raises on failure — callers decide whether to propagate or log."""
    import re
    from supabase import create_client
    sb = create_client(SUPABASE_URL, SUPABASE_KEY)

    from tools.db_modules.client import get_current_user_id
    user_id = get_current_user_id() or "shared"

    slug = f"{company}-{title}".lower()
    slug = re.sub(r"[^a-z0-9-]", "-", slug)[:50].strip("-")
    filename = Path(file_path).name
    storage_path = f"artifacts/users/{user_id}/applications/{slug}/{filename}"

    with open(file_path, "rb") as f:
        file_bytes = f.read()

    # upsert=true replaces any existing object at this path. Without it, a second
    # call for the same resume raises 409 and the DB points at the stale object.
    sb.storage.from_("artifacts").upload(
        storage_path,
        file_bytes,
        file_options={"upsert": "true"},
    )
    return storage_path


def update_pdf_path_in_db(job_id: str, pdf_path: str, docx_path: str = None, company: str = "", title: str = "") -> None:
    """Upload PDF (and DOCX, if present) to Storage and update the applications
    row to point at the newly-uploaded paths. Raises on any upload or DB failure
    — the caller (CLI or API) decides how to surface it."""
    try:
        relative_pdf = str(Path(pdf_path).relative_to(PROJECT_ROOT))
    except ValueError:
        relative_pdf = pdf_path

    from supabase import create_client
    sb = create_client(SUPABASE_URL, SUPABASE_KEY)

    update_data = {"resume_pdf_path": relative_pdf}

    pdf_storage_path = _upload_artifact_to_storage(job_id, pdf_path, company, title)
    update_data["resume_pdf_path_storage"] = pdf_storage_path

    if docx_path and Path(docx_path).exists():
        docx_storage_path = _upload_artifact_to_storage(job_id, docx_path, company, title)
        update_data["resume_docx_path"] = docx_storage_path

    res = sb.table("applications").update(update_data).eq("job_id", job_id).execute()
    if res.data:
        print(f"✅ Saved artifact paths to DB for job {job_id}")
    else:
        raise RuntimeError(f"applications row not found or not updated for job {job_id}")


def _safe_dirname(s: str) -> str:
    return re.sub(r"[^a-z0-9-]", "-", s.lower())[:40].strip("-")


# ─── CLI ──────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Generate a styled resume DOCX/PDF from markdown.")
    src = parser.add_mutually_exclusive_group(required=True)
    src.add_argument("--job-id",      help="Job UUID — fetches resume_md from DB")
    src.add_argument("--resume-path", help="Path to a local resume.md file")
    parser.add_argument("--output", "-o", help="Output path (.docx or .pdf)")
    parser.add_argument("--docx-only",  action="store_true", help="Generate DOCX only, skip PDF")
    parser.add_argument("--no-db-update", action="store_true", help="Skip writing PDF path to DB")
    args = parser.parse_args()

    company = ""
    title = ""
    if args.job_id:
        resume_md, company, title = fetch_resume_md_from_db(args.job_id)
        if args.output:
            output_path = args.output
        else:
            dir_name = f"{_safe_dirname(company)}-{_safe_dirname(title)}"
            out_dir  = PROJECT_ROOT / "output" / "applications" / dir_name
            out_dir.mkdir(parents=True, exist_ok=True)
            ext = ".docx" if args.docx_only else ".pdf"
            output_path = str(out_dir / f"resume{ext}")
    else:
        p = Path(args.resume_path)
        if not p.exists():
            print(f"❌ File not found: {args.resume_path}")
            sys.exit(1)
        resume_md = p.read_text(encoding="utf-8")
        ext = ".docx" if args.docx_only else ".pdf"
        output_path = args.output or str(p.with_suffix(ext))

    docx_path = None
    if args.docx_only or output_path.endswith(".docx"):
        docx_output = output_path if output_path.endswith(".docx") else output_path.replace(".pdf", ".docx")
        markdown_to_docx(resume_md, docx_output)
        docx_path = docx_output
    else:
        markdown_to_pdf(resume_md, output_path)
        # For PDF, also generate DOCX alongside it
        docx_path = output_path.replace(".pdf", ".docx")

    if args.job_id and not args.no_db_update and not args.docx_only:
        update_pdf_path_in_db(args.job_id, output_path, docx_path, company, title)


if __name__ == "__main__":
    main()
